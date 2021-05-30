import pymongo
from sshtunnel import SSHTunnelForwarder
import paramiko
import io
import base64
from flask_bcrypt import Bcrypt

with open('AIapp.pem', 'rb') as f:
    blob = base64.b64encode(f.read())

SSH_KEY_BLOB = blob
SSH_KEY_BLOB_DECODED = base64.b64decode(SSH_KEY_BLOB)
SSH_KEY = SSH_KEY_BLOB_DECODED.decode('utf-8')

pkey = paramiko.RSAKey.from_private_key(io.StringIO(SSH_KEY))

MONGO_HOST = "18.136.12.76"
MONGO_DB = "news_db"
MONGO_USER = "aiteam"
MONGO_PASS = "aiteam123"

server = SSHTunnelForwarder(
    MONGO_HOST,
    remote_bind_address=('127.0.0.1', 27017),
    ssh_username='ubuntu',
    ssh_pkey=pkey
)

server.start()
client = pymongo.MongoClient('127.0.0.1', server.local_bind_port, username=MONGO_USER, password=MONGO_PASS,
                             authSource='admin')
db = client[MONGO_DB]
users = db["users"]
bcrypt = Bcrypt(None)


def get_data(stock_code=0, limit=0, ascending=1):
    data = []
    if stock_code != 0:
        condition = {"_id": 0, "content": 0,
                     "classification_score": 0, "site": 0, "created_at": 0}
    else:
        condition = {"_id": 0, "content": 0,
                     "stock_code": 0, "classification_score": 0, "site": 0, "created_at": 0}

    if limit != 0:
        for obj in db.news_tbl.find({}, condition).sort('published_date', ascending).limit(limit):
            data.append(obj)
    else:
        for obj in db.news_tbl.find({}, condition):
            data.append(obj)
    return data


def get_daily_news(cat, s, e, l, asc=1):
    if cat is None:
        condition = {
            "category": {
                "$in": ["macro_news", "international_news", "stock_market"]
            },
            "published_date": {
                "$lt": e,
                "$gte": s
            }
        }
    else:
        condition = {
            "category": cat,
            "published_date": {
                "$lt": e,
                "$gte": s
            }
        }

    data = []

    if l != 0:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "stock_code": 0, "classification_score": 0, "site": 0,
                                     "created_at": 0}).sort('published_date', asc).limit(l):
            data.append(obj)
    else:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "stock_code": 0, "classification_score": 0, "site": 0,
                                     "created_at": 0}).sort('published_date', asc):
            data.append(obj)
    return data


def get_daily_news_main(cat, s, e, l, asc=1):
    if cat is None:
        condition = {
            "category": {
                "$in": ["macro_news", "international_news", "stock_market"]
            },
            "published_date": {
                "$lt": e,
                "$gte": s
            },
            "rank_score": {
                "$gte": 0.1
            }
        }
    else:
        condition = {
            "category": cat,
            "published_date": {
                "$lt": e,
                "$gte": s
            },
            "rank_score": {
                "$gte": 0.1
            }
        }

    data = []

    if l != 0:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "stock_code": 0, "classification_score": 0,
                                     "created_at": 0}).sort(
            'published_date', asc).limit(l):
            data.append(obj)
    else:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "stock_code": 0, "classification_score": 0,
                                     "created_at": 0}).sort(
            'published_date', asc):
            data.append(obj)
    return data


def get_stock_market_news(s, e, l, asc=1):
    condition = {
        "stock_code": {
            "$nin": ["---", None, ""]
        },
        "published_date": {
            "$lt": e,
            "$gte": s
        },
        "rank_score": {
            "$gte": 0
        }
    }

    data = []

    if l != 0:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "stock_code": 0, "classification_score": 0,
                                     "created_at": 0}).sort(
            'published_date', asc).limit(l):
            data.append(obj)
    else:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "stock_code": 0, "classification_score": 0,
                                     "created_at": 0}).sort(
            'published_date', asc):
            data.append(obj)
    return data


def get_data_by_category_and_date(cat, start, end, code, limit, ascending):
    data = []
    condition = {
        "category": cat,
        "published_date": {
            "$lt": end,
            "$gte": start
        },
        "$and": [
            {
                "stock_code": {
                    "$ne": ""
                }
            },
            {
                "stock_code": {
                    "$ne": None
                }
            }
        ]
    }

    if code != "All" and code != "-":
        condition["stock_code"] = code

    if limit != 0:
        for obj in db.news_tbl.find(condition, {"_id": 0, "content": 0, "classification_score": 0, "site": 0,
                                                "created_at": 0}).sort('published_date', ascending).limit(limit):
            data.append(obj)
    else:
        for obj in db.news_tbl.find(condition, {"_id": 0, "content": 0, "classification_score": 0, "site": 0,
                                                "created_at": 0}).sort('published_date', ascending):
            data.append(obj)
    return data


def get_data_by_content_and_keyword(keyword, start, end, limit, ascending):
    data = []
    condition = {
        "$or": [
            {
                "content": {
                    "$regex": keyword,
                    "$options": 'i'
                }
            },
            {
                "title": {
                    "$regex": keyword,
                    "$options": 'i'
                }
            }
        ],
        "published_date": {
            "$lt": end,
            "$gte": start
        }
    }

    if limit != 0:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "classification_score": 0, "site": 0,
                                     "created_at": 0}).sort('published_date', ascending).limit(limit):
            data.append(obj)
    else:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "classification_score": 0, "site": 0,
                                     "created_at": 0}).sort('published_date', ascending):
            data.append(obj)
    return data


def get_morning_news(s, e, l, asc=1):
    condition = {
        "category": {
            "$nin": ["macro_news", "international_news", "stock_market"]
        },
        "published_date": {
            "$lt": e,
            "$gte": s
        }
    }
    data = []

    if l != 0:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "classification_score": 0, "site": 0,
                                     "created_at": 0}).sort('published_date', asc).limit(l):
            data.append(obj)
    else:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "classification_score": 0, "site": 0,
                                     "created_at": 0}).sort('published_date', asc):
            data.append(obj)
    return data


def get_morning_news_main(s, e, l, asc=1):
    condition = {
        "category": {
            "$nin": ["macro_news", "international_news", "stock_market"]
        },
        "published_date": {
            "$lt": e,
            "$gte": s
        },
        "rank_score": {
            "$gte": 0.01
        }
    }
    data = []

    if l != 0:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "classification_score": 0, "created_at": 0}).sort(
            'published_date', asc).limit(l):
            data.append(obj)
    else:
        for obj in db.news_tbl.find(condition,
                                    {"_id": 0, "content": 0, "classification_score": 0, "created_at": 0}).sort(
            'published_date', asc):
            data.append(obj)
    return data


def update_news(condition, value):
    return db.news_tbl.update_one(condition, value)


import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from io import BytesIO
from bs4 import BeautifulSoup
import bs4


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def write_docx(daily_news, morning_brief, six_news, today):
    doc = docx.Document()

    h = doc.add_heading("IMPORTANT NEWS - {0}".format(today), 1)
    h.alignment = 1

    doc.add_heading("A. Daily news", 2)
    count = 0
    for _, news in daily_news.items():
        count += 1
        p = doc.add_paragraph("{0}. ".format(count))
        add_hyperlink(p, news['title'], news['source_url'])
        doc.add_paragraph(news['text_rank'])
        doc.add_paragraph("Nguồn {0}".format(news['site']))

    doc.add_heading("B. Morning brief", 2)
    count = 0
    for _, news in morning_brief.items():
        count += 1
        p = doc.add_paragraph("{0}. ".format(count))
        add_hyperlink(p, news['title'], news['source_url'])
        soup = BeautifulSoup(news['text_rank'])
        for t in soup.find("div", {"class": "content"}).contents:
            if type(t) != bs4.element.Tag:
                doc.add_paragraph(t)
            else:
                p = doc.add_paragraph("")
                add_hyperlink(p, t.find('a', href=True).string, t.find('a', href=True)['href'])

        doc.add_paragraph("Nguồn {0}".format(news['site']))

    doc.add_heading("C. Six news", 2)
    count = 0
    for _, news in six_news.items():
        count += 1
        p = doc.add_paragraph("{0}. ".format(count))
        add_hyperlink(p, news['title'], news['source_url'])
        doc.add_paragraph(news['brief_content'])
        doc.add_paragraph("Nguồn {0}".format(news['site']))

    in_memory_fp = BytesIO()
    doc.save(in_memory_fp)
    in_memory_fp.seek(0)
    return in_memory_fp

def get_max_rank(s_date, e_date, category):
    condition = {
        "category": category,
        "published_date": {
            "$lt": e_date,
            "$gte": s_date
        }
    }
    score = db.news_tbl.find_one(condition, sort=[('rank_score', -1)])
    print(score)
    return score['rank_score']